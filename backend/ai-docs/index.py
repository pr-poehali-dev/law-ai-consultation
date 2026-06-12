"""
AI-документы — генерация, анализ и продолжение документов.
Режимы: doc_generate, doc_continue, file_analyze, file_cleanup.
Таймаут: 120 секунд. Анализ: DeepSeek 2500 токенов.
"""
import json
import os
import re
import base64
import io
import time
import threading
import requests

from prompts import (
    SYSTEM_DOC_GENERATE, SYSTEM_FILE_ANALYZE_PROMPT, SYSTEM_FILE_QA_PROMPT,
    SYSTEM_DOC_BY_TYPE, REFUSAL_MARKERS, LEGAL_QUALITY_ADDON,
)
# doc_blocks.py — промты прямо в ai-docs, без зависимости от gigachat-proxy/prompts (пакет vs файл)
from doc_blocks import get_system_prompt, get_doc_label, BLOCK_BY_DOC_TYPE
_BLOCK_ROUTER_OK = True
print("[INIT] doc_blocks OK")
from state_duty import is_duty_query, get_duty_context_for_doc, DUTY_DOC_TYPES
from legal_docs_handler import get_legal_context_for_ai
from penalty_prompt import PENALTY_CALC_SYSTEM, PENALTY_CALC_PROMPT

YANDEX_MODEL = os.environ.get("YANDEX_MODEL_URI", "gpt://b1gd8kncmd8nf4j7h770/deepseek-v4-flash/latest")
YANDEX_MODEL_FAST = "gpt://b1gd8kncmd8nf4j7h770/deepseek-v4-flash/latest"
_IAM_TOKEN: str = os.environ.get("YANDEX_IAM_TOKEN", "").strip()

_http = requests.Session()
_http.headers.update({"Content-Type": "application/json"})

_RE_TRUNCATED = re.compile(r'[.!?»\d]\s*$')
_RE_PLACEHOLDER = re.compile(r'\{\{([^}]+)\}\}')
_RE_DOC_END_SPEECH = re.compile(r'(прошу\s+суд|прошу\s+уважаемый|на\s+основании\s+изложенного|итог|в\s+заключение)', re.I)
_RE_DOC_END_OTHER = re.compile(r'(подпись|реквизиты|экземпляр|дата\s*[:|]?\s*«|\d{1,2}\.\d{2}\.\d{4}|РЕШИЛ|ПОСТАНОВИЛ|обжаловано\s+в|вступает\s+в\s+законную\s+силу|апелляционном\s+порядке|ПРИНЯЛ|ПОСТАНОВЛЯЕТ|подписан[ао]|М\.П\.|место\s+печати|юридическ\w+\s+адрес|банковские\s+реквизиты|настоящий\s+договор\s+составлен|договор\s+вступает|стороны\s+пришли\s+к\s+соглашению)', re.I)

SPEECH_DOC_TYPES = {"court_speech"}

CORS = {
    "Access-Control-Allow-Origin": "*",
    "Access-Control-Allow-Methods": "GET, POST, OPTIONS",
    "Access-Control-Allow-Headers": "Content-Type, X-Auth-Token",
}

FILE_TTL = 1800
FILE_BUCKET = "files"
FILE_PREFIX = "temp-docs/"
MAX_FILE_MB = 10
ALLOWED_EXTS = {"pdf", "docx", "doc", "jpeg", "jpg", "png", "txt", "heic", "heif", "webp", "bmp", "tiff", "tif"}

# ── Персональные данные ────────────────────────────────────────────────────
_PD_PATTERNS = [
    (re.compile(r'\b\d{3}-\d{3}-\d{3}\s\d{2}\b'), '{{СНИЛС}}'),
    (re.compile(r'\bИНН[:\s]+\d{10,12}\b', re.IGNORECASE), '{{ИНН}}'),
    (re.compile(r'\bОГРН[:\s]+\d{13,15}\b', re.IGNORECASE), '{{ОГРН}}'),
    (re.compile(r'(паспорт|серия паспорта|выдан)[^,;\n]{0,40}\d{4}\s?\d{6}', re.IGNORECASE), '{{ПАСПОРТНЫЕ_ДАННЫЕ}}'),
    (re.compile(r'(?<!\d)\d{4}\s\d{6}(?!\d)'), '{{ПАСПОРТ}}'),
    (re.compile(r'(дата\s+рождения|д\.р\.|дата\sвыдачи|выдан\s+\d)[:\s]*\d{1,2}[./]\d{1,2}[./]\d{2,4}', re.IGNORECASE), '{{ДАТА}}'),
]

def _sanitize_doc_text(text: str) -> str:
    result = text
    for pattern, replacement in _PD_PATTERNS:
        result = pattern.sub(replacement, result)
    return result

# ── AI-вызовы ─────────────────────────────────────────────────────────────
MAX_HISTORY = 10

def is_refusal(text) -> bool:
    if not text:
        return False
    low = text.lower()
    return any(m in low for m in REFUSAL_MARKERS)

def call_yandex(system_prompt: str, messages: list, max_tokens: int = 1200, fast: bool = False, temperature: float = 0.3) -> str:
    recent = messages[-MAX_HISTORY:] if len(messages) > MAX_HISTORY else messages
    openai_messages = [{"role": "system", "content": system_prompt}] + [
        {"role": "user" if m.get("role") == "user" else "assistant",
         "content": m.get("content", m.get("text", ""))}
        for m in recent
    ]
    model = YANDEX_MODEL_FAST if fast else YANDEX_MODEL
    timeout = 30 if fast else 80
    resp = _http.post(
        "https://llm.api.cloud.yandex.net/v1/chat/completions",
        headers={"Authorization": f"Api-Key {_IAM_TOKEN}"},
        json={"model": model, "messages": openai_messages, "max_tokens": max_tokens, "temperature": temperature, "stream": False},
        timeout=timeout,
    )
    resp.raise_for_status()
    return resp.json()["choices"][0]["message"]["content"]

def call_deepseek(system_prompt: str, messages: list, max_tokens: int = 800, temperature: float = 0.3, timeout: int = 80) -> tuple:
    recent = messages[-MAX_HISTORY:] if len(messages) > MAX_HISTORY else messages
    openai_messages = [{"role": "system", "content": system_prompt}] + [
        {"role": "user" if m.get("role") == "user" else "assistant",
         "content": m.get("content", m.get("text", ""))}
        for m in recent
    ]
    resp = _http.post(
        "https://llm.api.cloud.yandex.net/v1/chat/completions",
        headers={"Authorization": f"Api-Key {_IAM_TOKEN}"},
        json={"model": YANDEX_MODEL, "messages": openai_messages, "max_tokens": max_tokens, "temperature": temperature, "stream": False},
        timeout=timeout,
    )
    resp.raise_for_status()
    resp_json = resp.json()
    choice = resp_json["choices"][0]
    text = choice["message"]["content"] or ""
    was_cut = choice.get("finish_reason") == "length"
    # Логируем реальный usage токенов для диагностики
    usage = resp_json.get("usage", {})
    print(f"[DEEPSEEK] finish={choice.get('finish_reason')} prompt_tokens={usage.get('prompt_tokens')} completion_tokens={usage.get('completion_tokens')} симв={len(text)}")
    return text, was_cut

# ── Файловые утилиты ───────────────────────────────────────────────────────
def get_s3():
    import boto3
    return boto3.client(
        "s3",
        endpoint_url="https://bucket.poehali.dev",
        aws_access_key_id=os.environ["AWS_ACCESS_KEY_ID"],
        aws_secret_access_key=os.environ["AWS_SECRET_ACCESS_KEY"],
    )

def save_temp_file(s3, data: bytes, filename: str, content_type: str) -> str:
    ts = int(time.time())
    key = f"{FILE_PREFIX}{ts}_{filename}"
    s3.put_object(Bucket=FILE_BUCKET, Key=key, Body=data, ContentType=content_type, Metadata={"uploaded_at": str(ts)})
    return key

def cleanup_temp_files(s3) -> list:
    deleted = []
    now = int(time.time())
    try:
        paginator = s3.get_paginator("list_objects_v2")
        for page in paginator.paginate(Bucket=FILE_BUCKET, Prefix=FILE_PREFIX):
            for obj in page.get("Contents", []):
                key = obj["Key"]
                basename = key.replace(FILE_PREFIX, "")
                parts = basename.split("_", 1)
                try:
                    uploaded_at = int(parts[0])
                    if now - uploaded_at >= FILE_TTL:
                        s3.delete_object(Bucket=FILE_BUCKET, Key=key)
                        deleted.append(key)
                except (ValueError, IndexError):
                    pass
    except Exception:
        pass
    return deleted

def extract_pdf_text(data: bytes, char_limit: int = 8000) -> str:
    import PyPDF2
    reader = PyPDF2.PdfReader(io.BytesIO(data))
    parts = []
    total = 0
    for page in reader.pages[:15]:
        text = page.extract_text() or ""
        parts.append(text)
        total += len(text)
        if total >= char_limit:
            break
    result = "\n".join(parts).strip()[:char_limit]
    if len(result.strip()) < 50 and _IAM_TOKEN:
        try:
            b64_pdf = base64.b64encode(data).decode("utf-8")
            vision_resp = _http.post(
                "https://vision.api.cloud.yandex.net/vision/v1/batchAnalyze",
                headers={"Authorization": f"Api-Key {_IAM_TOKEN}"},
                json={"analyzeSpecs": [{"content": b64_pdf, "features": [{"type": "TEXT_DETECTION", "textDetectionConfig": {"languageCodes": ["ru", "en"]}}], "mimeType": "application/pdf"}]},
                timeout=15,
            )
            if vision_resp.ok:
                blocks = (vision_resp.json().get("results", [{}])[0].get("results", [{}])[0].get("textDetection", {}).get("pages", []))
                lines = []
                for page in blocks:
                    for block in page.get("blocks", []):
                        for line in block.get("lines", []):
                            words = [w.get("text", "") for w in line.get("words", [])]
                            if words:
                                lines.append(" ".join(words))
                ocr_text = "\n".join(lines).strip()
                if ocr_text and len(ocr_text.strip()) > 30:
                    return ocr_text[:char_limit]
        except Exception:
            pass
    return result

def extract_docx_text(data: bytes, char_limit: int = 12000) -> str:
    """Читает .docx (python-docx) или .doc (Word 97-2003 через встроенный zipfile/regex)."""
    # Шаг 1: пробуем как .docx (OOXML = zip-архив с word/document.xml)
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(data))
        parts = []
        total = 0
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                parts.append(t)
                total += len(t)
                if total >= char_limit:
                    break
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t and total < char_limit:
                        parts.append(t)
                        total += len(t)
        result = "\n".join(parts)[:char_limit]
        if len(result.strip()) >= 30:
            return result
    except Exception:
        pass

    # Шаг 2: fallback — читаем как zip (docx без расширения или повреждённый)
    try:
        import zipfile
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            if "word/document.xml" in zf.namelist():
                xml_bytes = zf.read("word/document.xml")
                # Убираем все XML-теги, оставляем текст
                text = xml_bytes.decode("utf-8", errors="ignore")
                text = re.sub(r'<[^>]+>', ' ', text)
                text = re.sub(r'\s+', ' ', text).strip()
                if len(text) >= 30:
                    return text[:char_limit]
    except Exception:
        pass

    # Шаг 3: .doc (Word 97-2003 OLE) — парсим бинарник без внешних зависимостей
    # Сигнатура OLE: D0 CF 11 E0 A1 B1 1A E1
    if len(data) > 8 and data[:8] == b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1':
        try:
            # Ищем текстовые блоки UTF-16-LE (Word хранит текст как Unicode)
            decoded = data.decode("utf-16-le", errors="ignore")
            # Оставляем только печатные символы, убираем мусор
            cleaned = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', ' ', decoded)
            # Выбираем «слова» длиннее 2 символов
            words = re.findall(r'[А-ЯЁа-яёA-Za-z0-9][А-ЯЁа-яёA-Za-z0-9\s\.,;:!?\-«»"\']{2,}', cleaned)
            result = " ".join(words)[:char_limit]
            if len(result.strip()) >= 50:
                return result
        except Exception:
            pass
        try:
            # Попытка cp1251 (кириллица в старых doc)
            decoded = data.decode("cp1251", errors="ignore")
            cleaned = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', ' ', decoded)
            words = re.findall(r'[А-ЯЁа-яёA-Za-z0-9][А-ЯЁа-яёA-Za-z0-9\s\.,;:!?\-«»"\']{2,}', cleaned)
            result = " ".join(words)[:char_limit]
            if len(result.strip()) >= 50:
                return result
        except Exception:
            pass

    return ""

def _compress_image(image_data: bytes, max_bytes: int = 900_000) -> bytes:
    if len(image_data) <= max_bytes:
        return image_data
    try:
        from PIL import Image as PILImage
        img = PILImage.open(io.BytesIO(image_data))
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        max_side = 2000
        w, h = img.size
        if max(w, h) > max_side:
            ratio = max_side / max(w, h)
            img = img.resize((int(w * ratio), int(h * ratio)), PILImage.LANCZOS)
        buf = io.BytesIO()
        quality = 85
        while quality >= 40:
            buf.seek(0); buf.truncate()
            img.save(buf, format="JPEG", quality=quality, optimize=True)
            if buf.tell() <= max_bytes:
                break
            quality -= 15
        return buf.getvalue()
    except Exception:
        return image_data[:max_bytes]

def extract_image_text_ocr(image_data: bytes, ext: str) -> str:
    if not _IAM_TOKEN:
        return ""
    ext_lower = ext.lower()
    if ext_lower not in ("jpg", "jpeg", "png"):
        try:
            from PIL import Image as PILImage
            img = PILImage.open(io.BytesIO(image_data))
            if img.mode not in ("RGB", "L"):
                img = img.convert("RGB")
            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=85)
            image_data = buf.getvalue()
            ext_lower = "jpg"
        except Exception:
            pass
    mime_map = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png"}
    mime_type = mime_map.get(ext_lower, "image/jpeg")
    compressed = _compress_image(image_data)
    b64_image = base64.b64encode(compressed).decode("utf-8")
    try:
        resp = _http.post(
            "https://vision.api.cloud.yandex.net/vision/v1/batchAnalyze",
            headers={"Authorization": f"Api-Key {_IAM_TOKEN}"},
            json={"analyzeSpecs": [{"content": b64_image, "features": [{"type": "TEXT_DETECTION", "textDetectionConfig": {"languageCodes": ["ru", "en"]}}], "mimeType": mime_type}]},
            timeout=12,
        )
        if not resp.ok:
            return ""
        result = resp.json()
        blocks = (result.get("results", [{}])[0].get("results", [{}])[0].get("textDetection", {}).get("pages", []))
        lines = []
        for page in blocks:
            for block in page.get("blocks", []):
                for line in block.get("lines", []):
                    words = [w.get("text", "") for w in line.get("words", [])]
                    if words:
                        lines.append(" ".join(words))
        return "\n".join(lines).strip()
    except Exception:
        return ""

def analyze_file_with_deepseek(text: str, comment: str, n_docs: int = 1) -> str:
    """Анализ документов через DeepSeek v3 (Yandex Cloud).

    Из логов: 900 completion_tokens = ~433 симв (кириллица ≈ 0.48 симв/токен).
    Для полного ответа (5 разделов, ~1200-1500 симв) нужно max_tokens=2500.
    timeout=95с — DeepSeek генерирует 2500 токенов за ~30-60с реально.
    """
    clean_text = _sanitize_doc_text(text)

    if comment:
        # Пользователь задал вопрос — отвечаем на него с опорой на документ
        user_content = (
            f"Вопрос пользователя: {comment}\n\n"
            f"{'Документы' if n_docs > 1 else 'Документ'} ({n_docs} шт.):\n\n{clean_text}"
        )
        system_prompt = SYSTEM_FILE_QA_PROMPT
    else:
        # Анализ без конкретного вопроса
        if n_docs > 1:
            note = (
                f"\n\nВАЖНО: загружено {n_docs} документов. "
                f"Раздел 3 (Риски) — укажи риски по каждому документу отдельно. "
                f"Раздел 4 (Рекомендации) — учти все документы вместе."
            )
        else:
            note = ""
        user_content = (
            f"{'Документы' if n_docs > 1 else 'Документ'} для анализа "
            f"({'каждый проанализируй отдельно' if n_docs > 1 else ''}):\n\n{clean_text}"
        )
        system_prompt = SYSTEM_FILE_ANALYZE_PROMPT + note

    result, was_cut = call_deepseek(
        system_prompt,
        [{"role": "user", "content": user_content}],
        max_tokens=2500,   # 2500 токенов × 0.48 симв/токен ≈ 1200 симв — полный ответ
        temperature=0.15,
        timeout=95,        # 95с < 120с таймаут функции, хватает на 2500 токенов
    )
    print(f"[FILE_ANALYZE] {'WARN was_cut' if was_cut else 'OK'}: симв={len(result)}, docs={n_docs}")
    return result.strip()


def handler(event: dict, context) -> dict:
    """AI-документы: генерация, анализ файлов, продолжение. Таймаут 90с."""
    if event.get("httpMethod") == "OPTIONS":
        return {"statusCode": 200, "headers": CORS, "body": ""}
    if event.get("httpMethod") == "GET":
        return {"statusCode": 200, "headers": {**CORS, "Content-Type": "application/json"},
                "body": json.dumps({"ok": True, "service": "ai-docs"})}

    headers = event.get("headers") or {}
    token = headers.get("X-Auth-Token") or headers.get("x-auth-token", "")

    body = {}
    if event.get("body"):
        try:
            body = json.loads(event["body"])
        except Exception:
            pass

    mode = body.get("mode", "")

    try:
        # ── Очистка временных файлов ─────────────────────────────────────────
        if mode == "file_cleanup":
            s3 = get_s3()
            deleted = cleanup_temp_files(s3)
            return {"statusCode": 200, "headers": {**CORS, "Content-Type": "application/json"},
                    "body": json.dumps({"deleted": len(deleted), "keys": deleted}, ensure_ascii=False)}

        # ── Продолжение обрезанного документа ───────────────────────────────
        if mode == "doc_continue":
            doc_type = body.get("doc_type", "claim")
            partial = body.get("partial", "").strip()
            if not partial:
                return {"statusCode": 400, "headers": CORS, "body": json.dumps({"error": "partial required"})}
            sp_cont = get_system_prompt(doc_type, LEGAL_QUALITY_ADDON)
            system_prompt = sp_cont if sp_cont else SYSTEM_DOC_BY_TYPE.get(doc_type, SYSTEM_DOC_GENERATE)
            context_tail = partial[-800:]
            prompt = (
                f"Документ был обрезан. Продолжи с того места, где остановился, без повторения уже написанного.\n\n"
                f"Конец уже написанного текста:\n...{context_tail}\n\n"
                f"Продолжай документ до финального блока с реквизитами, подписями и датой. "
                f"Незаполненные поля — метки {{{{ПОЛЕ_НАЗВАНИЕ}}}}."
            )
            answer = call_yandex(system_prompt, [{"role": "user", "content": prompt}], max_tokens=3500, temperature=0.15)
            truncated = not bool(re.search(r'(подпись|реквизиты|экземпляр|дата\s*[:|]?\s*«|\d{1,2}\.\d{2}\.\d{4})', answer[-300:], re.I))
            placeholders = list(dict.fromkeys(re.findall(r'\{\{([^}]+)\}\}', answer)))
            return {"statusCode": 200, "headers": {**CORS, "Content-Type": "application/json"},
                    "body": json.dumps({"answer": answer, "placeholders": placeholders, "truncated": truncated}, ensure_ascii=False)}

        # ── Генерация документа ──────────────────────────────────────────────
        if mode == "doc_generate":
            _mode_start = time.time()
            if not token:
                return {"statusCode": 401, "headers": {**CORS, "Content-Type": "application/json"},
                        "body": json.dumps({"error": "Требуется авторизация"}, ensure_ascii=False)}
            doc_type = body.get("doc_type", "claim")
            details = body.get("details", "").strip()
            file_b64 = body.get("file", "")
            filename = body.get("filename", "")
            if not details:
                return {"statusCode": 400, "headers": CORS, "body": json.dumps({"error": "details required"})}

            label = get_doc_label(doc_type)
            sp = get_system_prompt(doc_type, LEGAL_QUALITY_ADDON)
            system_prompt = sp if sp else SYSTEM_DOC_BY_TYPE.get(doc_type, SYSTEM_DOC_GENERATE)
            print(f"[DOC_GEN] system_prompt len={len(system_prompt)} doc_type={doc_type} has_block={'YES' if sp else 'FALLBACK'}")
            if doc_type in DUTY_DOC_TYPES:
                system_prompt = system_prompt + (
                    "\n\n[ВАЖНО О ГОСПОШЛИНЕ] Используй ТОЛЬКО ставки из правовой базы. "
                    "НЕ выдумывай суммы из памяти."
                )

            # ── Извлечение текста из файлов (старый формат: file+filename, новый: files[]) ──
            file_context = ""

            # Новый формат: files = [{b64, name}, ...] до 3 шт
            files_list = body.get("files", [])
            # Обратная совместимость: одиночный file
            if not files_list and file_b64 and filename:
                files_list = [{"b64": file_b64, "name": filename}]

            if files_list:
                extracted_parts = []
                for idx, fobj in enumerate(files_list[:3]):
                    fb64 = fobj.get("b64", "")
                    fname = fobj.get("name", f"file{idx+1}")
                    if not fb64:
                        continue
                    try:
                        file_data = base64.b64decode(fb64)
                        ext = fname.rsplit(".", 1)[-1].lower() if "." in fname else ""
                        # Лимит 1800 симв на файл — не перегружаем промт, сохраняем токены генерации
                        if ext == "pdf":
                            text = extract_pdf_text(file_data)[:1800]
                        elif ext in ("docx", "doc"):
                            text = extract_docx_text(file_data)[:1800]
                        elif ext in ("jpg", "jpeg", "png"):
                            text = extract_image_text_ocr(file_data, ext)[:1400]
                        else:
                            text = ""
                        if text.strip():
                            # Без квадратных скобок — иначе модель путает с маркерами блоков документа
                            extracted_parts.append(f"Документ {idx+1} ({fname}):\n{text.strip()}")
                            print(f"[DOC_GEN] файл {idx+1} ({fname}): извлечено {len(text)} симв")
                    except Exception as ex:
                        print(f"[DOC_GEN] файл {idx+1} ({fname}) ошибка: {ex}")

                if extracted_parts:
                    file_context = "\n\n".join(extracted_parts)

            chat_history = body.get("chat_history", [])
            history_context = ""
            if chat_history:
                last_pairs = chat_history[-5:]
                history_context = "КОНТЕКСТ ИЗ ПРЕДЫДУЩЕЙ КОНСУЛЬТАЦИИ (используй все упомянутые факты, стороны, суммы, даты):\n"
                for msg in last_pairs:
                    role_label = "Пользователь" if msg.get("role") == "user" else "Юрист"
                    content_raw = msg.get("content", "")[:800]
                    result = content_raw
                    for pattern, replacement in _PD_PATTERNS:
                        result = pattern.sub(replacement, result)
                    history_context += f"{role_label}: {result}\n"
                history_context += "\nВАЖНО: подстави все конкретные данные из диалога напрямую в документ. Где данных нет — используй метки {{{{ПОЛЕ_НАЗВАНИЕ}}}}.\n\n"

            speech_style = ""
            if doc_type in SPEECH_DOC_TYPES:
                speech_style = (
                    "ОРАТОРСКИЙ СТИЛЬ (обязательно): короткие предложения (7–12 слов), паузы (//), "
                    "тройные перечисления, сильные глаголы, один тезис на абзац. "
                    "Ключевую мысль — в начало и в конец. Без воды и длинных оборотов.\n\n"
                )

            # ── Типы документов по длине — ДОЛЖНЫ БЫТЬ ОБЪЯВЛЕНЫ ДО THREADS ──────
            _SHORT_DOC_TYPES = {
                "contract_receipt", "special_pd_consent", "special_medical_consent",
                "labor_hire_app", "labor_quit_app", "labor_vacation_app",
                "labor_vacation_exit", "labor_quit_withdraw",
            }
            _MEDIUM_DOC_TYPES = {
                "labor_order_hire", "labor_order_dismiss",
                "labor_order_bonus", "labor_order_discipline",
                "petition_postpone", "petition_absence", "petition_attach",
                "application_writ", "application_clarify",
            }
            _LONG_DOC_TYPES = {
                # Исковые и жалобы
                "claim", "claim_debt", "claim_divorce", "claim_property",
                "claim_consumer", "claim_damage", "claim_ownership",
                "claim_paternity", "claim_eviction", "claim_alimony",
                "claim_admin", "claim_order", "claim_counter", "claim_interim",
                "appeal", "cassation", "supervisory", "partial_appeal",
                "criminal_cassation", "criminal_appeal",
                # Судебные решения и определения
                "court_decision", "court_ruling", "court_order",
                "decision_court", "decision_sole", "decision_meeting",
                "decision_board", "decision_liquidation",
                # Отзывы, возражения
                "response_to_claim", "objection_appeal", "objection_cassation",
                "written_explanations",
                # Уголовный процесс
                "criminal_exclude_evidence", "criminal_125", "criminal_terminate",
                "criminal_measure", "criminal_attach_evidence", "criminal_statement",
                "criminal_witness", "criminal_special_order", "criminal_reinvestigate",
                # Договоры — все сложные
                "corporate_charter", "corporate_collective", "corporate_rules",
                "corporate_founding", "corporate_branch", "corporate_job_desc",
                "contract_gov", "contract_service_state",
                "website_terms", "website_privacy", "website_eula",
                "website_offer", "website_aup", "website_disclaimer",
                "contract_marriage", "special_will", "special_inheritance_contract",
                "labor_contract", "contract_sale", "contract_rent", "contract_work",
                "contract_services", "contract_loan", "contract_license",
                "contract_cession", "contract_partnership",
                "contract_children", "contract_property_split",
                "contract_preliminary", "contract_mediation",
                # Трудовые сложные
                "labor_addendum", "labor_termination",
                "labor_downtime_notice", "labor_dismiss_notice", "labor_layoff_notice",
                # Корпоративные
                "corporate_sole_decision", "corporate_meeting", "corporate_board",
                "corporate_salary", "corporate_bonus", "corporate_pd", "corporate_secret",
                "corporate_staffing", "corporate_accounting",
            }

            # Параллельно запрашиваем все 4 категории правовой базы
            _duty_db_result: list = []
            _case_law_result: list = []
            _definitions_result: list = []
            _codex_result: list = []

            def _fetch_duty_db():
                if doc_type in DUTY_DOC_TYPES:
                    _duty_db_result.append(get_legal_context_for_ai("state_duty", max_files=2, max_chars=3000, query=details))
            # Для коротких кадровых документов правовая база не нужна — не перегружаем промт
            def _fetch_case_law():
                if doc_type not in (_MEDIUM_DOC_TYPES | _SHORT_DOC_TYPES):
                    _case_law_result.append(get_legal_context_for_ai("case_law", max_files=2, max_chars=3000, query=details))
            def _fetch_definitions():
                if doc_type not in (_MEDIUM_DOC_TYPES | _SHORT_DOC_TYPES):
                    _definitions_result.append(get_legal_context_for_ai("court_definitions", max_files=2, max_chars=3000, query=details))
            def _fetch_codex():
                if doc_type not in (_MEDIUM_DOC_TYPES | _SHORT_DOC_TYPES):
                    _codex_result.append(get_legal_context_for_ai("codex", max_files=3, max_chars=3000, query=details))

            threads = [
                threading.Thread(target=_fetch_duty_db, daemon=True),
                threading.Thread(target=_fetch_case_law, daemon=True),
                threading.Thread(target=_fetch_definitions, daemon=True),
                threading.Thread(target=_fetch_codex, daemon=True),
            ]
            for t in threads: t.start()
            for t in threads: t.join(timeout=8)

            # ГОСПОШЛИНА: берём ТОЛЬКО из правовой базы (раздел state_duty)
            # Запрещено выдумывать ставки или брать их из памяти модели
            duty_db = _duty_db_result[0] if _duty_db_result else ""
            duty_block = ""
            if doc_type in DUTY_DOC_TYPES:
                if duty_db:
                    duty_block = (
                        "\n\n[ОБЯЗАТЕЛЬНО] ГОСПОШЛИНА — используй ТОЛЬКО ставки и льготы из правовой базы ниже.\n"
                        "ЗАПРЕЩЕНО выдумывать суммы или применять знания из памяти — только эти данные:\n"
                        + duty_db
                        + "\n\nПРАВИЛА:\n"
                        "1. Если в правовой базе есть подходящая ставка — укажи её точно.\n"
                        "2. Если есть льгота по этой категории дел — укажи льготу (без расчёта пошлины).\n"
                        "3. Если данных в базе недостаточно — вставь метку {{ЦЕНА_ИСКА}} и формулу из базы.\n"
                        "4. Не вводи цифры которых нет в правовой базе.\n"
                    )
                else:
                    # Нет данных в БД — не указываем госпошлину вообще
                    duty_block = (
                        "\n\n[ВАЖНО] По госпошлине данных в правовой базе нет — НЕ указывай конкретные суммы.\n"
                        "Используй метку {{ГОСПОШЛИНА}} или формулировку 'согласно НК РФ ст. 333.19/333.21'.\n"
                    )
            case_law_block = _case_law_result[0] if _case_law_result else ""
            definitions_block = _definitions_result[0] if _definitions_result else ""
            codex_block = _codex_result[0] if _codex_result else ""
            extra_context = duty_block + case_law_block + definitions_block + codex_block
            print(f"[DOC_GEN] Правовая база: duty={bool(duty_block)}, case_law={bool(case_law_block)}, definitions={bool(definitions_block)}, codex={bool(codex_block)}")

            # Данные из файлов встраиваются прямо в описание ситуации — без маркеров-скобок
            file_note = ""
            if file_context:
                file_note = (
                    "\n\nДополнительные данные из прикреплённых документов "
                    "(используй ФИО, адреса, суммы, даты, реквизиты напрямую — не ставь заглушки там где данные есть):\n"
                    f"{file_context}"
                )

            prompt = (
                history_context + speech_style
                + f"Составь {label} на основании следующего описания ситуации:\n\n{details}"
                + file_note + "\n\n"
                + LEGAL_QUALITY_ADDON + extra_context
                + f"\nТам где не хватает конкретных данных (ФИО, адрес, номер дела и т.д.) — "
                f"используй метки-заглушки {{{{ПОЛЕ_НАЗВАНИЕ}}}} (русский язык, подчёркивание). "
                f"Запрещены [...] и ___.\n\n"
                f"ОБЯЗАТЕЛЬНО: структурируй документ блоками — каждый маркер ([ШАПКА], [ЗАГОЛОВОК], [ТЕЛО], [ТРЕБОВАНИЯ], [ПРИЛОЖЕНИЯ], [ПОДПИСЬ]) на ОТДЕЛЬНОЙ строке. "
                f"НЕ добавляй блоки [ОБОСНОВАНИЕ] и [ПРИМЕЧАНИЯ] — они не нужны."
            )
            # DeepSeek fallback использует тот же prompt что и Яндекс
            raw_prompt = prompt

            # Лимиты токенов (типы уже объявлены выше перед threads)
            _EXTRA_LONG_DOC_TYPES = {
                # Решения суда — самые объёмные, могут требовать 7000-8000 токенов
                "court_decision", "decision_court",
                # Уставы и корпоративные регламенты
                "corporate_charter", "corporate_collective", "corporate_rules",
                # Сложные договоры с большим числом условий
                "contract_marriage", "contract_gov", "contract_service_state",
                "special_inheritance_contract", "special_will",
                "website_terms", "website_privacy", "website_eula",
            }
            if doc_type in _SHORT_DOC_TYPES:
                _max_tokens = 600
            elif doc_type in _MEDIUM_DOC_TYPES:
                _max_tokens = 2000
            elif doc_type in _EXTRA_LONG_DOC_TYPES:
                _max_tokens = 8000
            elif doc_type in _LONG_DOC_TYPES:
                _max_tokens = 6000
            else:
                _max_tokens = 6000

            # Если есть файлы — промт длиннее, ограничиваем его чтобы не съел токены генерации
            has_files = bool(file_context)
            if has_files:
                # Обрезаем до 20000 симв — оставляем место для генерации (DeepSeek 128k контекст)
                MAX_PROMPT_CHARS = 20000
                if len(raw_prompt) > MAX_PROMPT_CHARS:
                    # Обрезаем extra_context (правовая база) — менее критично чем файл
                    overflow = len(raw_prompt) - MAX_PROMPT_CHARS
                    if extra_context and len(extra_context) > overflow + 500:
                        extra_context = extra_context[:len(extra_context) - overflow - 500] + "\n..."
                    elif file_note and len(file_note) > overflow + 200:
                        file_note = file_note[:len(file_note) - overflow - 200] + "\n..."
                    raw_prompt = (
                            history_context + speech_style
                            + f"Составь {label} на основании следующего описания ситуации:\n\n{details}"
                            + file_note + "\n\n"
                            + LEGAL_QUALITY_ADDON + extra_context
                            + f"\nТам где не хватает конкретных данных (ФИО, адрес, номер дела и т.д.) — "
                            f"используй метки-заглушки {{{{ПОЛЕ_НАЗВАНИЕ}}}} (русский язык, подчёркивание). "
                            f"Запрещены [...] и ___."
                        )
                print(f"[DOC_GEN] промт с файлами: {len(raw_prompt)} симв")

            # Генерация только через DeepSeek (deepseek-v32 на Yandex Cloud)
            # Таймаут функции должен быть 120с — DeepSeek берёт 60-90с на большой документ
            answer = ""
            try:
                answer, was_cut = call_deepseek(
                    system_prompt,
                    [{"role": "user", "content": raw_prompt}],
                    max_tokens=_max_tokens,
                    temperature=0.15,
                    timeout=110,  # 110с < таймаут функции 120с
                )
                print(f"[DOC_GEN] DeepSeek OK симв={len(answer)} was_cut={was_cut} has_files={has_files}")
            except Exception as e:
                print(f"[DOC_GEN] DeepSeek упал: {e}")

            if not answer:
                return {"statusCode": 500, "headers": {**CORS, "Content-Type": "application/json"},
                        "body": json.dumps({"error": "Не удалось сгенерировать документ. Попробуйте ещё раз."}, ensure_ascii=False)}

            if doc_type in SPEECH_DOC_TYPES:
                truncated = not bool(_RE_DOC_END_SPEECH.search(answer[-400:]))
            else:
                truncated = not bool(_RE_DOC_END_OTHER.search(answer[-300:]))
            placeholders = list(dict.fromkeys(_RE_PLACEHOLDER.findall(answer)))
            # Рекомендации теперь запрашиваются отдельно (mode=doc_recommendations) после показа документа
            return {"statusCode": 200, "headers": {**CORS, "Content-Type": "application/json"},
                    "body": json.dumps({"answer": answer, "placeholders": placeholders, "truncated": truncated}, ensure_ascii=False)}

        # ── Анализ файлов ────────────────────────────────────────────────────
        if mode == "file_analyze":
            comment = body.get("comment", "").strip()
            iam_token = os.environ["YANDEX_IAM_TOKEN"].strip()
            mime_map = {
                "pdf": "application/pdf", "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "doc": "application/msword", "jpeg": "image/jpeg", "jpg": "image/jpeg",
                "png": "image/png", "txt": "text/plain",
            }

            raw_files = body.get("files", [])
            if not raw_files:
                f_b64 = body.get("file", "")
                f_name = body.get("filename", "document")
                if f_b64:
                    raw_files = [{"file": f_b64, "filename": f_name}]
            if not raw_files:
                return {"statusCode": 400, "headers": {**CORS, "Content-Type": "application/json"},
                        "body": json.dumps({"error": "file required"}, ensure_ascii=False)}
            if len(raw_files) > 3:
                raw_files = raw_files[:3]

            extract_results = [None] * len(raw_files)

            # Бюджет текста для DeepSeek: суммарно ~8000 симв
            n_raw = len(raw_files)
            if n_raw == 1:
                per_file_char_limit = 8000   # 1×8000
            elif n_raw == 2:
                per_file_char_limit = 4000   # 2×4000=8000
            else:
                per_file_char_limit = 2700   # 3×2700=8100

            def _extract_one(idx, fi):
                fb64 = fi.get("file", "")
                fname = fi.get("filename", "document")
                if not fb64:
                    return
                fdata = base64.b64decode(fb64)
                if len(fdata) > MAX_FILE_MB * 1024 * 1024:
                    return
                ext = fname.rsplit(".", 1)[-1].lower() if "." in fname else ""
                if ext not in ALLOWED_EXTS:
                    return
                ct = mime_map.get(ext, "application/octet-stream")
                def _bg_up(d=fdata, n=fname, c=ct):
                    try:
                        s3 = get_s3(); save_temp_file(s3, d, n, c); cleanup_temp_files(s3)
                    except Exception: pass
                threading.Thread(target=_bg_up, daemon=True).start()
                if ext == "pdf":
                    extract_results[idx] = (fname, extract_pdf_text(fdata, char_limit=per_file_char_limit), None)
                elif ext in ("docx", "doc"):
                    extract_results[idx] = (fname, extract_docx_text(fdata, char_limit=per_file_char_limit), None)
                elif ext == "txt":
                    extract_results[idx] = (fname, fdata.decode("utf-8", errors="replace")[:per_file_char_limit], None)
                else:
                    compressed = _compress_image(fdata, max_bytes=700_000)
                    ocr_b64 = base64.b64encode(compressed).decode("utf-8")
                    t = extract_image_text_ocr(compressed, ext)
                    if not t or len(t.strip()) < 15:
                        extract_results[idx] = (fname, "", ocr_b64)
                    else:
                        extract_results[idx] = (fname, t[:per_file_char_limit], None)

            extract_threads = [threading.Thread(target=_extract_one, args=(i, fi), daemon=True) for i, fi in enumerate(raw_files)]
            for t in extract_threads: t.start()
            # PDF/DOCX: ~1–2с, OCR-фото: ~12с, все потоки параллельно → достаточно 20с
            for t in extract_threads: t.join(timeout=20)

            file_texts = [r for r in extract_results if r is not None]
            filenames_list = [r[0] for r in file_texts]
            first_ocr_b64 = next((r[2] for r in file_texts if r[2]), None)

            if not file_texts:
                return {"statusCode": 400, "headers": {**CORS, "Content-Type": "application/json"},
                        "body": json.dumps({"error": "Не удалось обработать файлы. Проверьте формат и размер."}, ensure_ascii=False)}

            combined_parts = []
            all_ocr_failed = True
            for fname, txt, ocr_b64 in file_texts:
                if txt and len(txt.strip()) >= 15:
                    all_ocr_failed = False
                    label = f"[Файл: {fname}]" if len(file_texts) > 1 else ""
                    combined_parts.append(f"{label}\n{txt.strip()}" if label else txt.strip())

            combined_text = "\n\n---\n\n".join(combined_parts)
            combined_filename = filenames_list[0] if len(filenames_list) == 1 else f"{filenames_list[0]} +{len(filenames_list)-1}"

            # Vision fallback для фото без OCR-текста
            if all_ocr_failed and first_ocr_b64:
                try:
                    n_files = len(file_texts)
                    vision_intro = f"Ты — опытный юрист РФ. Пользователь загрузил {'фото документа' if n_files == 1 else f'{n_files} фото документов'}.\n\n"
                    vision_task = (
                        f"Вопрос пользователя: {comment}\n\nОтвечай СТРОГО на этот вопрос, используя документ как источник фактов. Не делай общий анализ — только ответ со ссылками на статьи законов РФ."
                        if comment else
                        "СТРОГО ЗАПРЕЩЕНО: пересказывать текст документа.\n\nДай краткое практическое юридическое заключение:\n**Тип и суть** — одной фразой.\n**Правовые риски** — со статьёй закона РФ: 🔴 критично / 🟡 существенно / 🟢 незначительно.\n**Что делать** — 1–3 практических шага."
                    )
                    user_msg_content = [
                        {"type": "text", "text": vision_intro + vision_task + "\n\nЕсли фото нечёткое — скажи об этом."},
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{first_ocr_b64}"}},
                    ]
                    vision_resp = requests.post(
                        "https://llm.api.cloud.yandex.net/v1/chat/completions",
                        headers={"Authorization": f"Api-Key {iam_token}", "Content-Type": "application/json"},
                        json={"model": "gpt://b1g2k5n3ojr7ik7lv73l/yandexgpt-vision-lite/latest", "messages": [{"role": "user", "content": user_msg_content}], "max_tokens": 2000, "temperature": 0.1},
                        timeout=30,
                    )
                    if vision_resp.ok:
                        answer = vision_resp.json()["choices"][0]["message"]["content"]
                        return {"statusCode": 200, "headers": {**CORS, "Content-Type": "application/json"},
                                "body": json.dumps({"answer": answer, "filename": filenames_list[0], "delete_at": int(time.time()) + FILE_TTL}, ensure_ascii=False)}
                except Exception:
                    pass

            if not combined_text or len(combined_text.strip()) < 20:
                return {"statusCode": 400, "headers": {**CORS, "Content-Type": "application/json"},
                        "body": json.dumps({"error": "Не удалось извлечь текст из файлов. Попробуйте PDF или DOCX."}, ensure_ascii=False)}

            analysis_result = [None]
            analysis_error = [None]
            hint_result = [None]
            n_files = len(file_texts)

            def _do_analysis():
                try:
                    analysis_result[0] = analyze_file_with_deepseek(combined_text, comment, n_docs=n_files)
                except Exception as _ae:
                    analysis_error[0] = str(_ae)
                    print(f"[FILE_ANALYZE] Поток упал: {_ae}")

            def _do_hint():
                doc_types_list = "claim — исковое заявление, pretension — досудебная претензия, complaint — жалоба, application — заявление/ходатайство, notification — уведомление, order — приказ, contract — договор ГПХ, business_contract — коммерческий договор, court_speech — судебная речь, response_to_claim — отзыв на исковое заявление, objection — возражение, appeal — апелляционная жалоба, cassation — кассационная жалоба, supervisory — надзорная жалоба"
                doc_intro = f"Пользователь загрузил {n_files} документа" if n_files > 1 else "Пользователь загрузил документ"
                if comment:
                    hint_prompt = (
                        f"Ты — помощник юриста. {doc_intro} и ЯВНО УКАЗАЛ какой документ нужно составить.\n\n"
                        f"ЗАПРОС ПОЛЬЗОВАТЕЛЯ (ПРИОРИТЕТ): {comment}\n\n"
                        f"Доступные типы документов: {doc_types_list}\n\n"
                        f"Текст документа:\n{combined_text[:3000]}\n\n"
                        f"Ответь СТРОГО в JSON:\n"
                        f'{{"doc_type": "id_типа", "details": "подробное описание для генерации", "doc_label": "название на русском"}}'
                    )
                else:
                    hint_prompt = (
                        f"Ты — помощник юриста. {doc_intro}. Определи какой ответный документ нужно составить.\n"
                        f"Доступные типы: {doc_types_list}\n\nТекст:\n{combined_text[:5000]}\n\n"
                        f"Ответь СТРОГО в JSON:\n"
                        f'{{"doc_type": "id_типа", "details": "подробное описание", "doc_label": "название на русском"}}'
                    )
                try:
                    resp = requests.post(
                        "https://llm.api.cloud.yandex.net/v1/chat/completions",
                        headers={"Authorization": f"Api-Key {iam_token}", "Content-Type": "application/json"},
                        json={"model": YANDEX_MODEL_FAST, "messages": [{"role": "user", "content": hint_prompt}], "max_tokens": 800, "temperature": 0.1, "stream": False},
                        timeout=30,
                    )
                    resp.raise_for_status()
                    raw = resp.json()["choices"][0]["message"]["content"]
                    match = re.search(r'\{[\s\S]*\}', raw)
                    if match:
                        hint_result[0] = json.loads(match.group())
                except Exception:
                    pass

            _analyze_start = time.time()
            t_analysis = threading.Thread(target=_do_analysis, daemon=True)
            t_hint = threading.Thread(target=_do_hint, daemon=True)
            # hint запускаем первым — он быстрый (YandexGPT fast, 30с)
            # analysis запускается параллельно — DeepSeek 2500 токенов, до 90с
            t_hint.start(); t_analysis.start()
            # Бюджет: извлечение 20с + анализ DeepSeek 95с = 115с < 120с таймаут функции
            t_analysis.join(timeout=97)   # чуть больше timeout DeepSeek=95с чтобы поток успел вернуть результат
            t_hint.join(timeout=2)        # hint уже давно завершён (~5-15с), join мгновенный

            elapsed = time.time() - _analyze_start
            print(f"[FILE_ANALYZE] analysis elapsed={elapsed:.1f}s, result={'OK' if analysis_result[0] else 'FAIL'}")

            if not analysis_result[0]:
                err_msg = (
                    f"Анализ занял слишком много времени ({elapsed:.0f}с). Попробуйте с меньшим числом файлов."
                    if not analysis_error[0] else
                    f"Не удалось проанализировать документ: {analysis_error[0][:120]}"
                )
                return {"statusCode": 502, "headers": {**CORS, "Content-Type": "application/json"},
                        "body": json.dumps({"error": err_msg}, ensure_ascii=False)}

            response_data = {
                "answer": analysis_result[0],
                "filename": combined_filename,
                "delete_at": int(time.time()) + FILE_TTL,
                "extracted_text": combined_text[:6000],
            }
            if hint_result[0]:
                response_data["doc_hint"] = hint_result[0]
            return {"statusCode": 200, "headers": {**CORS, "Content-Type": "application/json"},
                    "body": json.dumps(response_data, ensure_ascii=False)}

        # ── Расчёт неустойки ─────────────────────────────────────────────────
        if mode == "penalty_calc":
            if not token:
                return {"statusCode": 401, "headers": {**CORS, "Content-Type": "application/json"},
                        "body": json.dumps({"error": "Unauthorized"}, ensure_ascii=False)}
            calc_data = body.get("calc_data", "")
            if not calc_data:
                return {"statusCode": 400, "headers": {**CORS, "Content-Type": "application/json"},
                        "body": json.dumps({"error": "calc_data required"}, ensure_ascii=False)}
            # Используем replace вместо format — промпт содержит {плейсхолдеры} для AI, не для Python
            penalty_prompt_text = PENALTY_CALC_PROMPT.replace("{calc_data}", calc_data)
            pen_answer = ""
            # DeepSeek — точнее считает неустойки по ГК РФ
            try:
                pen_answer, _ = call_deepseek(PENALTY_CALC_SYSTEM, [{"role": "user", "content": penalty_prompt_text}], max_tokens=2500, temperature=0.05, timeout=75)
                print(f"[PENALTY_CALC] DeepSeek ответил, len={len(pen_answer)}")
            except Exception as e:
                print(f"[PENALTY_CALC] DeepSeek упал: {e} → fallback YandexGPT")
                try:
                    pen_answer = call_yandex(PENALTY_CALC_SYSTEM, [{"role": "user", "content": penalty_prompt_text}], max_tokens=2000, fast=True, temperature=0.05)
                except Exception as e2:
                    print(f"[PENALTY_CALC] YandexGPT тоже упал: {e2}")
            if not pen_answer:
                return {"statusCode": 500, "headers": {**CORS, "Content-Type": "application/json"},
                        "body": json.dumps({"error": "Не удалось выполнить расчёт. Попробуйте ещё раз."}, ensure_ascii=False)}
            return {"statusCode": 200, "headers": {**CORS, "Content-Type": "application/json"},
                    "body": json.dumps({"answer": pen_answer}, ensure_ascii=False)}

        # ── Генерация документа по рекомендации ──────────────────────────────
        if mode == "rec_doc_generate":
            if not token:
                return {"statusCode": 401, "headers": {**CORS, "Content-Type": "application/json"},
                        "body": json.dumps({"error": "Unauthorized"}, ensure_ascii=False)}
            rec_doc_type = body.get("rec_doc_type", "")
            rec_context = body.get("rec_context", "")
            rec_reason = body.get("rec_reason", "")
            if not rec_doc_type or not rec_context:
                return {"statusCode": 400, "headers": {**CORS, "Content-Type": "application/json"},
                        "body": json.dumps({"error": "rec_doc_type and rec_context required"}, ensure_ascii=False)}

            REC_DOC_LABELS = {
                "motion_restore_term": "ходатайство о восстановлении срока",
                "motion_evidence": "ходатайство об истребовании доказательств",
                "motion_witness": "ходатайство о вызове свидетеля",
                "motion_third_party": "ходатайство о привлечении третьего лица",
                "motion_expertise": "ходатайство о назначении экспертизы",
                "motion_enforcement": "ходатайство об обеспечении иска",
                "pretension": "досудебную претензию",
                "complaint": "жалобу",
                "appeal": "апелляционную жалобу",
            }
            rec_label = REC_DOC_LABELS.get(rec_doc_type, "дополнительный документ")
            rec_system = (
                "Ты — опытный российский юрист. Составь юридически грамотный документ. "
                "Используй формат с блоками [ШАПКА], [ЗАГОЛОВОК], [ТЕКСТ], [ПРИЛОЖЕНИЯ], [ПОДПИСЬ]. "
                "Где нет конкретных данных — ставь {{ПОЛЕ_НАЗВАНИЕ}}. "
                "Запрещены [...] и ___."
            )
            rec_prompt_text = (
                f"Составь {rec_label} на основании следующего контекста.\n\n"
                f"Контекст основного документа:\n{rec_context[:2000]}\n\n"
                f"Обоснование необходимости: {rec_reason}\n\n"
                "Составь полный готовый документ со ссылками на нормы ТК/ГПК/АПК/ГК РФ."
            )
            rec_answer = ""
            try:
                # YandexGPT (не fast) — timeout=80с, max_tokens=2800 безопасно
                rec_answer = call_yandex(rec_system, [{"role": "user", "content": rec_prompt_text}], max_tokens=2800, temperature=0.1)
                print(f"[REC_DOC] YandexGPT OK len={len(rec_answer)}")
            except Exception as e:
                print(f"[REC_DOC] Яндекс упал: {e} → DeepSeek fallback")
            if not rec_answer or is_refusal(rec_answer):
                try:
                    rec_answer, _ = call_deepseek(rec_system, [{"role": "user", "content": rec_prompt_text}], max_tokens=2500, temperature=0.1, timeout=65)
                except Exception as e:
                    print(f"[REC_DOC] DeepSeek упал: {e}")
            if not rec_answer:
                return {"statusCode": 500, "headers": {**CORS, "Content-Type": "application/json"},
                        "body": json.dumps({"error": "Не удалось создать документ. Попробуйте ещё раз."}, ensure_ascii=False)}
            rec_placeholders = list(dict.fromkeys(_RE_PLACEHOLDER.findall(rec_answer)))
            return {"statusCode": 200, "headers": {**CORS, "Content-Type": "application/json"},
                    "body": json.dumps({"answer": rec_answer, "placeholders": rec_placeholders, "label": rec_label}, ensure_ascii=False)}

        # ── Анализ рекомендаций (вызывается отдельно после показа документа) ──
        if mode == "doc_recommendations":
            if not token:
                return {"statusCode": 401, "headers": {**CORS, "Content-Type": "application/json"},
                        "body": json.dumps({"error": "Unauthorized"}, ensure_ascii=False)}
            doc_name = body.get("doc_name", "Документ")
            doc_content = body.get("doc_content", "").strip()
            if not doc_content:
                return {"statusCode": 400, "headers": {**CORS, "Content-Type": "application/json"},
                        "body": json.dumps({"error": "doc_content required"}, ensure_ascii=False)}
            # Подтягиваем актуальную судебную практику для рекомендаций
            legal_ctx = ""
            try:
                legal_ctx = get_legal_context_for_ai("case_law", max_files=2, max_chars=2000, query=doc_name)
            except Exception:
                pass
            rec_sys = (
                "Ты — опытный юрист РФ. Анализируй документ и давай ценные рекомендации пользователю. "
                "Рекомендации должны быть конкретными и полезными. "
                "Отвечай ТОЛЬКО корректным JSON без пояснений вне JSON."
            )
            rec_pr = (
                f"Документ «{doc_name}»:\n{doc_content[:2500]}\n\n"
                + (f"Актуальная судебная практика по теме:\n{legal_ctx}\n\n" if legal_ctx else "")
                + "Дай рекомендации пользователю. Типы рекомендаций:\n"
                '1. {"type": "general", "title": "Краткий заголовок", "reason": "...", "advice": "Конкретный совет"} — общие юридические советы, на что обратить внимание\n'
                '2. {"type": "state_duty", "title": "Госпошлина", "reason": "...", "duty_note": "Замечание"} — ТОЛЬКО если в документе явно неверно указана или пропущена госпошлина\n'
                '3. {"type": "penalty_calc", "title": "Расчёт неустойки", "reason": "..."} — ТОЛЬКО если документ подразумевает взыскание неустойки/пени и расчёта нет\n'
                '4. {"type": "doc", "doc_type": "TYPE", "title": "...", "reason": "..."} — ТОЛЬКО если реально необходим дополнительный документ\n'
                "   doc_type: motion_restore_term | motion_evidence | motion_witness | motion_third_party | motion_expertise | motion_enforcement | pretension | complaint | appeal\n\n"
                "ПРАВИЛА:\n"
                "— general-советы давай всегда (1-2 штуки) если есть что посоветовать\n"
                "— state_duty — только при реальной ошибке в госпошлине\n"
                "— penalty_calc и doc — только при реальной необходимости\n"
                "— Максимум 3 рекомендации. Пустой список если совсем нечего добавить.\n\n"
                'Верни: {"recommendations": [...]}'
            )
            recs_raw = []
            try:
                raw = call_yandex(rec_sys, [{"role": "user", "content": rec_pr}], max_tokens=700, fast=True, temperature=0.15)
                m = re.search(r'\{[\s\S]*\}', raw)
                if m:
                    parsed = json.loads(m.group())
                    recs_raw = parsed.get("recommendations", [])
                    if not isinstance(recs_raw, list): recs_raw = []
                    recs_raw = recs_raw[:3]
                    print(f"[DOC_RECS] got {len(recs_raw)} recs")
            except Exception as e:
                print(f"[DOC_RECS] ошибка: {e}")
            return {"statusCode": 200, "headers": {**CORS, "Content-Type": "application/json"},
                    "body": json.dumps({"recommendations": recs_raw}, ensure_ascii=False)}

        # ── AI-анализ документа (помощник) ────────────────────────────────────
        if mode == "doc_ai_review":
            if not token:
                return {"statusCode": 401, "headers": {**CORS, "Content-Type": "application/json"},
                        "body": json.dumps({"error": "Unauthorized"}, ensure_ascii=False)}
            doc_name = body.get("doc_name", "Документ")
            doc_content = body.get("doc_content", "").strip()
            if not doc_content:
                return {"statusCode": 400, "headers": {**CORS, "Content-Type": "application/json"},
                        "body": json.dumps({"error": "doc_content required"}, ensure_ascii=False)}
            review_system = (
                "Ты — опытный юрист РФ. Анализируй юридические документы кратко и ёмко. "
                "Структурируй ответ чётко по разделам с эмодзи-маркерами. "
                "Отвечай строго в указанном JSON-формате."
            )
            # Обрезаем документ до 2500 символов — оптимально для анализа
            doc_for_review = doc_content[:2500]
            review_prompt = (
                f"Проанализируй документ «{doc_name}» (первые символы):\n\n"
                f"{doc_for_review}\n\n"
                "Дай заключение СТРОГО в формате (используй эти маркеры):\n"
                "⚖️ ЮРИДИЧЕСКАЯ КОРРЕКТНОСТЬ\n[1-2 предложения: есть ли ошибки/неточности]\n\n"
                "📊 ПЕРСПЕКТИВА\n[высокая/средняя/низкая + одна фраза обоснования]\n\n"
                "📋 РЕКОМЕНДАЦИИ\n[2-3 конкретных улучшения]\n\n"
                "Верни JSON:\n"
                '{"answer": "текст заключения выше", "recommendations": ['
                '{"type": "penalty_calc", "title": "Расчёт неустойки", "reason": "..."} | '
                '{"type": "doc", "doc_type": "TYPE", "title": "...", "reason": "..."}]}\n'
                "doc_type: motion_restore_term|motion_evidence|motion_witness|motion_third_party|"
                "motion_expertise|motion_enforcement|pretension|complaint|appeal\n"
                "Максимум 2 рекомендации. Пустой список если не нужно."
            )
            review_answer = ""
            review_recs = []
            # Лимит 2000 токенов — быстро и в таймаут
            REVIEW_MAX_TOKENS = 2000
            try:
                raw = call_yandex(review_system, [{"role": "user", "content": review_prompt}], max_tokens=REVIEW_MAX_TOKENS, fast=True, temperature=0.1)
                m = re.search(r'\{[\s\S]*\}', raw)
                if m:
                    parsed = json.loads(m.group())
                    review_answer = parsed.get("answer", raw)
                    review_recs = parsed.get("recommendations", [])
                else:
                    review_answer = raw
            except Exception as e:
                print(f"[DOC_REVIEW] Яндекс упал: {e}")
            if not review_answer:
                return {"statusCode": 500, "headers": {**CORS, "Content-Type": "application/json"},
                        "body": json.dumps({"error": "Не удалось провести анализ. Попробуйте ещё раз."}, ensure_ascii=False)}
            return {"statusCode": 200, "headers": {**CORS, "Content-Type": "application/json"},
                    "body": json.dumps({"answer": review_answer, "recommendations": review_recs}, ensure_ascii=False)}

        # ── Редактирование документа AI-помощником ────────────────────────────
        if mode == "doc_edit":
            if not token:
                return {"statusCode": 401, "headers": {**CORS, "Content-Type": "application/json"},
                        "body": json.dumps({"error": "Unauthorized"}, ensure_ascii=False)}
            doc_name = body.get("doc_name", "Документ")
            doc_content = body.get("doc_content", "").strip()
            edit_instruction = body.get("edit_instruction", "").strip()
            if not doc_content or not edit_instruction:
                return {"statusCode": 400, "headers": {**CORS, "Content-Type": "application/json"},
                        "body": json.dumps({"error": "doc_content and edit_instruction required"}, ensure_ascii=False)}

            doc_len = len(doc_content)

            # Судебную практику подтягиваем ТОЛЬКО если пользователь прямо просит
            extra_edit_ctx = ""
            LEGAL_KEYWORDS = [
                "практик", "прецедент", "суды решают", "судебн", "аналогичн",
                "пошлин", "госпошлин", "333.19", "333.21",
            ]
            if any(k in edit_instruction.lower() for k in LEGAL_KEYWORDS):
                try:
                    extra_edit_ctx = get_legal_context_for_ai("case_law", max_files=1, max_chars=1000, query=edit_instruction)
                except Exception:
                    pass

            # ── Умная стратегия точечного редактирования ─────────────────────
            # Шаг 1: находим релевантный фрагмент документа (~2000 символов)
            # Шаг 2: AI редактирует ТОЛЬКО этот фрагмент
            # Шаг 3: вставляем отредактированный фрагмент обратно в документ
            # Это позволяет не нарушать структуру документа и экономить токены

            CONTEXT_WINDOW = 2500  # символов вокруг найденного места

            # Ищем подходящий фрагмент по ключевым словам из инструкции
            def _find_edit_fragment(content: str, instruction: str) -> tuple[int, int]:
                """Возвращает (start, end) позиции фрагмента для редактирования."""
                # Нормализуем инструкцию для поиска
                instr_lower = instruction.lower()
                # Разбиваем на строки, ищем наиболее релевантную
                lines = content.split("\n")
                best_line_idx = 0
                best_score = 0
                # Ключевые слова из инструкции (убираем стоп-слова)
                stop = {"в", "и", "на", "по", "с", "к", "из", "для", "что", "как", "это", "а", "но", "или", "не"}
                keywords = [w for w in instr_lower.split() if len(w) > 2 and w not in stop][:10]
                for i, line in enumerate(lines):
                    line_lower = line.lower()
                    score = sum(1 for kw in keywords if kw in line_lower)
                    if score > best_score:
                        best_score = score
                        best_line_idx = i
                # Находим позицию выбранной строки в исходном тексте
                char_pos = sum(len(lines[j]) + 1 for j in range(best_line_idx))
                # Берём фрагмент вокруг найденного места
                start = max(0, char_pos - CONTEXT_WINDOW // 2)
                end = min(len(content), char_pos + CONTEXT_WINDOW // 2)
                # Расширяем до границ строк
                while start > 0 and content[start - 1] != "\n":
                    start -= 1
                while end < len(content) and content[end] != "\n":
                    end += 1
                return start, end

            # Для коротких документов — передаём целиком
            if doc_len <= CONTEXT_WINDOW * 2:
                doc_fragment = doc_content
                frag_start, frag_end = 0, doc_len
                is_full_doc = True
            else:
                frag_start, frag_end = _find_edit_fragment(doc_content, edit_instruction)
                doc_fragment = doc_content[frag_start:frag_end]
                is_full_doc = False

            edit_system = (
                "Ты — юрист-редактор РФ. Вносишь точечные правки в юридический документ.\n"
                "ПРАВИЛА (обязательны):\n"
                "1. Изменяй ТОЛЬКО то, что прямо указано в инструкции пользователя.\n"
                "2. Весь остальной текст — ДОСЛОВНО, без малейших изменений.\n"
                "3. Не добавляй заголовки типа 'ДОКУМЕНТ «название»:' — только сам текст.\n"
                "4. Верни ТОЛЬКО отредактированный текст — без предисловий и пояснений.\n"
                "5. После текста добавь: ##CHANGES## и кратко — что именно изменено."
            )

            edit_prompt = (
                f"ФРАГМЕНТ ДОКУМЕНТА ДЛЯ РЕДАКТИРОВАНИЯ:\n"
                "─────────────────────\n"
                f"{doc_fragment}\n"
                "─────────────────────\n\n"
                f"ЗАДАЧА РЕДАКТОРА: {edit_instruction}\n"
                + (f"\n{extra_edit_ctx}\n" if extra_edit_ctx else "")
                + "\nВерни отредактированный фрагмент без лишних слов, затем ##CHANGES## и что изменено."
            )

            # DeepSeek — основной (лучшее качество редактирования, конфиденциальность)
            # YandexGPT fast — резервный (на случай таймаута DeepSeek)
            # Лимит токенов: 2500 — достаточно для полного документа, укладываемся в 90с
            MAX_TOKENS = 2500
            edit_answer = ""

            try:
                edit_answer, was_cut = call_deepseek(
                    edit_system,
                    [{"role": "user", "content": edit_prompt}],
                    max_tokens=MAX_TOKENS,
                    temperature=0.05,
                    timeout=65,
                )
                print(f"[DOC_EDIT] DeepSeek OK was_cut={was_cut} len={len(edit_answer)} doc={doc_len}")
            except Exception as e:
                print(f"[DOC_EDIT] DeepSeek упал: {e} → YandexGPT fallback")
                try:
                    edit_answer = call_yandex(
                        edit_system,
                        [{"role": "user", "content": edit_prompt}],
                        max_tokens=MAX_TOKENS,
                        fast=True,
                        temperature=0.05,
                    )
                    print(f"[DOC_EDIT] YandexGPT OK len={len(edit_answer)}")
                except Exception as e2:
                    print(f"[DOC_EDIT] YandexGPT тоже упал: {e2}")

            if not edit_answer:
                return {"statusCode": 500, "headers": {**CORS, "Content-Type": "application/json"},
                        "body": json.dumps({"error": "Не удалось отредактировать документ. Попробуйте ещё раз."}, ensure_ascii=False)}

            # Извлекаем список изменений
            changes_summary = ""
            if "##CHANGES##" in edit_answer:
                parts_ch = edit_answer.split("##CHANGES##", 1)
                edit_answer = parts_ch[0].strip()
                changes_summary = parts_ch[1].strip() if len(parts_ch) > 1 else ""

            # Собираем полный документ: до фрагмента + отредактированный + после фрагмента
            if is_full_doc:
                final_doc = edit_answer
            else:
                before = doc_content[:frag_start]
                after = doc_content[frag_end:]
                final_doc = (before + "\n" + edit_answer + "\n" + after).strip()
                print(f"[DOC_EDIT] Собран: before={len(before)} fragment={len(edit_answer)} after={len(after)} total={len(final_doc)}")

            return {"statusCode": 200, "headers": {**CORS, "Content-Type": "application/json"},
                    "body": json.dumps({
                        "answer": final_doc,
                        "partial_note": "",
                        "changes_summary": changes_summary,
                    }, ensure_ascii=False)}

        return {"statusCode": 400, "headers": {**CORS, "Content-Type": "application/json"},
                "body": json.dumps({"error": f"Unknown mode: {mode}"})}

    except Exception as e:
        if hasattr(e, "response") and e.response is not None:
            code = e.response.status_code
            try:
                detail = e.response.json()
            except Exception:
                detail = e.response.text[:300]
            return {"statusCode": 502, "headers": CORS,
                    "body": json.dumps({"error": f"HTTP {code}: {detail}"}, ensure_ascii=False)}
        return {"statusCode": 500, "headers": CORS, "body": json.dumps({"error": str(e)})}